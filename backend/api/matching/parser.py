"""Resume document parsing and structural analysis.

Handles PDF and DOCX, extracts the whole document (the previous implementation
capped at three pages), and detects the structure that ATS systems rely on.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

# Magic bytes, checked server-side. Extension alone is not evidence of format.
PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"  # DOCX is a zip container.

MAX_CHARS = 200_000

SECTION_PATTERNS: dict[str, re.Pattern[str]] = {
    "contact": re.compile(r"^\s*(contact|personal\s+details?|personal\s+info)", re.I | re.M),
    "summary": re.compile(r"^\s*(summary|profile|objective|about\s+me)\b", re.I | re.M),
    "experience": re.compile(
        r"^\s*(work\s+)?(experience|employment|professional\s+experience|career\s+history)\b",
        re.I | re.M,
    ),
    "education": re.compile(r"^\s*(education|academic|qualifications?)\b", re.I | re.M),
    "skills": re.compile(r"^\s*(skills?|technical\s+skills?|competenc)", re.I | re.M),
    "projects": re.compile(r"^\s*(projects?|portfolio)\b", re.I | re.M),
    "certifications": re.compile(r"^\s*(certifications?|licen[sc]es?|awards?)\b", re.I | re.M),
}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?){2,4}\d{2,4}")
URL_RE = re.compile(r"(?:https?://|www\.)[^\s)]+", re.I)
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w-]+", re.I)
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
BULLET_RE = re.compile(r"^\s*[•▪◦●\-\*•]\s+", re.M)

# Verbs that read as accomplishments rather than duties.
ACTION_VERBS = frozenset(
    """
    achieved administered analyzed architected automated built collaborated conducted
    consolidated constructed created cut decreased delivered demonstrated designed
    developed devised directed drove earned eliminated engineered enhanced established
    exceeded executed expanded facilitated forecast founded generated grew headed
    implemented improved increased influenced initiated innovated instituted integrated
    introduced launched led maintained managed maximized mentored migrated minimized
    modernized negotiated optimized orchestrated organized overhauled oversaw performed
    pioneered planned prioritized produced programmed proposed rebuilt reduced refactored
    resolved restructured revamped saved scaled shipped simplified spearheaded
    standardized streamlined strengthened supervised supported surpassed transformed
    trained translated upgraded won
    """.split()
)


class DocumentParseError(Exception):
    """Raised when a document cannot be read as text."""


@dataclass
class ParsedDocument:
    text: str = ""
    page_count: int = 0
    word_count: int = 0
    sections: list[str] = field(default_factory=list)
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    urls: list[str] = field(default_factory=list)
    has_linkedin: bool = False
    bullet_count: int = 0
    action_verb_count: int = 0
    years_mentioned: list[int] = field(default_factory=list)


def sniff_format(data: bytes) -> str:
    """Return ``"pdf"`` or ``"docx"`` from the file's magic bytes."""
    if data.startswith(PDF_MAGIC):
        return "pdf"
    if data.startswith(ZIP_MAGIC):
        return "docx"
    raise DocumentParseError(
        "Unrecognised file format. Upload a PDF or DOCX resume."
    )


def _extract_pdf(data: bytes) -> tuple[str, int]:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # Many resumes are "encrypted" with an empty owner password.
            try:
                reader.decrypt("")
            except Exception as exc:  # pragma: no cover - pypdf variance
                raise DocumentParseError("This PDF is password protected.") from exc
        pages = [page.extract_text() or "" for page in reader.pages]
    except DocumentParseError:
        raise
    except (PdfReadError, Exception) as exc:
        raise DocumentParseError(f"Could not read the PDF: {exc}") from exc
    return "\n".join(pages), len(pages)


def _extract_docx(data: bytes) -> tuple[str, int]:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseError(f"Could not read the DOCX file: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts), 0


def extract_text(data: bytes) -> tuple[str, int]:
    """Extract raw text and page count from a resume file."""
    fmt = sniff_format(data)
    text, pages = _extract_pdf(data) if fmt == "pdf" else _extract_docx(data)
    text = text[:MAX_CHARS]
    if not text.strip():
        raise DocumentParseError(
            "No text could be extracted. The file may be a scanned image; "
            "export a text-based PDF instead."
        )
    return text, pages


def analyze(text: str, page_count: int = 0) -> ParsedDocument:
    """Derive structural signals from already-extracted resume text."""
    words = text.split()
    lowered = {w.strip(".,;:()").lower() for w in words}

    # finditer, not findall: the pattern has groups, so findall would return
    # tuples of group captures rather than whole matches.
    phones = [m.group(0).strip() for m in PHONE_RE.finditer(text)]
    phones = [p for p in phones if sum(c.isdigit() for c in p) >= 9]

    return ParsedDocument(
        text=text,
        page_count=page_count,
        word_count=len(words),
        sections=[name for name, pattern in SECTION_PATTERNS.items() if pattern.search(text)],
        emails=sorted(set(EMAIL_RE.findall(text))),
        phones=sorted(set(phones))[:3],
        urls=sorted(set(URL_RE.findall(text)))[:10],
        has_linkedin=bool(LINKEDIN_RE.search(text)),
        bullet_count=len(BULLET_RE.findall(text)),
        action_verb_count=len(lowered & ACTION_VERBS),
        years_mentioned=sorted({int(m.group(0)) for m in YEAR_RE.finditer(text)}),
    )


def parse_document(data: bytes) -> ParsedDocument:
    """Full pipeline: bytes -> extracted text -> structural analysis."""
    text, pages = extract_text(data)
    return analyze(text, pages)
